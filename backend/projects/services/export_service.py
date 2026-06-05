"""
Export Service

Generates content pack exports in multiple formats:
  - JSON (immediate, no dependencies)
  - DOCX (requires python-docx)
  - PDF (requires reportlab)

All formats include:
  hooks, titles, captions, scripts, ctas, hashtags,
  transcript, summary, keywords, topics, moments
"""

import json
import logging
import io
from datetime import datetime

logger = logging.getLogger(__name__)


def build_content_pack(project) -> dict:
    """
    Assemble a complete content pack dict from project assets and intelligence.
    """
    from projects.models import GeneratedAsset, ContentIntelligenceRecord, Moment, TranscriptRecord

    def _assets_of_type(asset_type):
        return list(
            GeneratedAsset.objects.filter(project=project, type=asset_type)
            .values_list('content', flat=True)
        )

    # Hashtags
    hashtag_content = GeneratedAsset.objects.filter(
        project=project, type='HASHTAG'
    ).values_list('content', flat=True).first() or ''
    hashtags = [h.strip() for h in hashtag_content.split() if h.strip().startswith('#')]

    # Intelligence
    intelligence = {}
    try:
        intel = project.intelligence
        intelligence = {
            'summary': intel.summary,
            'topics': intel.topics,
            'keywords': intel.keywords,
            'entities': intel.entities,
            'viral_score': intel.viral_score,
        }
    except ContentIntelligenceRecord.DoesNotExist:
        pass

    # Transcript
    transcript_text = ''
    try:
        source = project.sources.filter(status='COMPLETED').first()
        if source:
            try:
                transcript_text = source.transcript.normalized_text
            except Exception:
                transcript_text = source.text_content
    except Exception:
        pass

    # Moments
    moments = list(
        Moment.objects.filter(project=project)
        .values('title', 'category', 'score', 'start_time', 'end_time', 'excerpt')
        .order_by('-score')[:20]
    )

    pack = {
        'project': {
            'id': project.id,
            'name': project.name,
            'description': project.description,
            'exported_at': datetime.utcnow().isoformat() + 'Z',
        },
        'intelligence': intelligence,
        'moments': moments,
        'hooks': _assets_of_type('HOOK'),
        'titles': _assets_of_type('TITLE'),
        'captions': _assets_of_type('CAPTION'),
        'ctas': _assets_of_type('CTA'),
        'scripts': _assets_of_type('SCRIPT'),
        'hashtags': hashtags,
        'transcript': transcript_text[:5000] if transcript_text else '',
    }
    return pack


def export_json(project) -> bytes:
    """Return JSON bytes of the full content pack."""
    pack = build_content_pack(project)
    return json.dumps(pack, indent=2, ensure_ascii=False).encode('utf-8')


def export_docx(project) -> bytes:
    """Return DOCX bytes of the content pack."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("[Export] python-docx not installed. Run: pip install python-docx")
        raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")

    pack = build_content_pack(project)
    doc = Document()

    # Title
    title_para = doc.add_heading(f"ViralOps Content Pack — {pack['project']['name']}", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Exported: {pack['project']['exported_at']}")
    doc.add_paragraph(f"Description: {pack['project']['description'] or 'N/A'}")

    # Intelligence
    if pack['intelligence']:
        doc.add_heading('Content Intelligence', level=1)
        intel = pack['intelligence']
        if intel.get('summary'):
            doc.add_heading('Summary', level=2)
            doc.add_paragraph(intel['summary'])
        if intel.get('topics'):
            doc.add_heading('Topics', level=2)
            doc.add_paragraph(' • '.join(intel['topics']))
        if intel.get('keywords'):
            doc.add_heading('Keywords', level=2)
            doc.add_paragraph(', '.join(intel['keywords']))
        if intel.get('viral_score') is not None:
            doc.add_paragraph(f"Viral Score: {intel['viral_score']}/100")

    # Moments
    if pack['moments']:
        doc.add_heading('AI Detected Moments', level=1)
        for m in pack['moments']:
            doc.add_heading(f"[{m['category']}] {m['title']} — Score: {m['score']}/100", level=2)
            if m.get('start_time') and m.get('end_time'):
                doc.add_paragraph(f"Time: {m['start_time']} → {m['end_time']}")
            if m.get('excerpt'):
                p = doc.add_paragraph(m['excerpt'])
                p.style = 'Quote'

    # Content assets
    sections = [
        ('Hooks', pack['hooks']),
        ('Titles', pack['titles']),
        ('Captions', pack['captions']),
        ('CTAs', pack['ctas']),
        ('Scripts', pack['scripts']),
    ]
    for section_name, items in sections:
        if items:
            doc.add_heading(section_name, level=1)
            for i, item in enumerate(items, 1):
                doc.add_heading(f"{section_name[:-1]} {i}", level=2)
                doc.add_paragraph(str(item))

    if pack['hashtags']:
        doc.add_heading('Hashtags', level=1)
        doc.add_paragraph(' '.join(pack['hashtags']))

    if pack.get('transcript'):
        doc.add_heading('Transcript (Preview)', level=1)
        doc.add_paragraph(pack['transcript'][:3000] + '...' if len(pack['transcript']) > 3000 else pack['transcript'])

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_pdf(project) -> bytes:
    """Return PDF bytes of the content pack."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_CENTER
    except ImportError:
        logger.error("[Export] reportlab not installed. Run: pip install reportlab")
        raise ImportError("reportlab is required for PDF export. Install with: pip install reportlab")

    pack = build_content_pack(project)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle', parent=styles['Title'],
        fontSize=20, textColor=colors.HexColor('#7C3AED'), spaceAfter=12
    )
    h1_style = ParagraphStyle(
        'CustomH1', parent=styles['Heading1'],
        fontSize=14, textColor=colors.HexColor('#7C3AED'), spaceBefore=16, spaceAfter=8
    )
    h2_style = ParagraphStyle(
        'CustomH2', parent=styles['Heading2'],
        fontSize=11, textColor=colors.HexColor('#6D28D9'), spaceBefore=10, spaceAfter=4
    )
    body_style = styles['Normal']
    quote_style = ParagraphStyle(
        'Quote', parent=styles['Normal'],
        leftIndent=20, rightIndent=20, fontName='Helvetica-Oblique',
        textColor=colors.HexColor('#6B7280'), spaceAfter=8
    )

    # Title
    story.append(Paragraph(f"ViralOps Content Pack", title_style))
    story.append(Paragraph(f"{pack['project']['name']}", h1_style))
    story.append(Paragraph(f"Exported: {pack['project']['exported_at']}", body_style))
    story.append(Spacer(1, 0.3*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#7C3AED')))
    story.append(Spacer(1, 0.3*cm))

    # Intelligence
    intel = pack.get('intelligence', {})
    if intel:
        story.append(Paragraph("Content Intelligence", h1_style))
        if intel.get('summary'):
            story.append(Paragraph("Summary", h2_style))
            story.append(Paragraph(intel['summary'], body_style))
        if intel.get('topics'):
            story.append(Paragraph("Topics", h2_style))
            story.append(Paragraph(' • '.join(intel['topics']), body_style))
        if intel.get('keywords'):
            story.append(Paragraph("Keywords", h2_style))
            story.append(Paragraph(', '.join(intel['keywords']), body_style))
        if intel.get('viral_score') is not None:
            story.append(Paragraph(f"Viral Score: {intel['viral_score']}/100", body_style))
        story.append(Spacer(1, 0.3*cm))

    # Moments
    moments = pack.get('moments', [])
    if moments:
        story.append(Paragraph("AI Detected Moments", h1_style))
        for m in moments[:10]:
            story.append(Paragraph(
                f"[{m['category']}] {m['title']} — Score: {m['score']}/100", h2_style
            ))
            if m.get('start_time') and m.get('end_time'):
                story.append(Paragraph(f"Time: {m['start_time']} → {m['end_time']}", body_style))
            if m.get('excerpt'):
                story.append(Paragraph(m['excerpt'][:300], quote_style))
        story.append(Spacer(1, 0.3*cm))

    # Content sections
    sections = [
        ('Hooks', pack.get('hooks', [])),
        ('Titles', pack.get('titles', [])),
        ('Captions', pack.get('captions', [])),
        ('CTAs', pack.get('ctas', [])),
        ('Scripts', pack.get('scripts', [])),
    ]
    for section_name, items in sections:
        if items:
            story.append(Paragraph(section_name, h1_style))
            for i, item in enumerate(items, 1):
                story.append(Paragraph(f"{section_name[:-1]} {i}", h2_style))
                story.append(Paragraph(str(item)[:500], body_style))
            story.append(Spacer(1, 0.2*cm))

    if pack.get('hashtags'):
        story.append(Paragraph("Hashtags", h1_style))
        story.append(Paragraph(' '.join(pack['hashtags']), body_style))

    doc.build(story)
    return buffer.getvalue()
